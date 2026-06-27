import uuid
import io
import qrcode
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage
from django.core.files.base import ContentFile


def log_action(user, action, model_name, object_id, details=None, ip=None):
    from apps.meetings.models import ActionLog
    ActionLog.objects.create(
        user=user,
        action=action,
        model_name=model_name,
        object_id=object_id,
        details=details or {},
        ip_address=ip,
    )


def generate_qr_code(data: str) -> bytes:
    qr = qrcode.QRCode(version=1, box_size=8, border=2)
    qr.add_data(data)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    buf = io.BytesIO()
    img.save(buf, format='PNG')
    return buf.getvalue()


def generate_pdf(meeting):
    from apps.meetings.models import Document

    qr_code_str = str(uuid.uuid4())[:12].upper()
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                             rightMargin=2*cm, leftMargin=2*cm,
                             topMargin=2*cm, bottomMargin=2*cm)

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('title', parent=styles['Heading1'],
                                  fontSize=18, spaceAfter=6, alignment=1,
                                  textColor=colors.HexColor('#1a5276'))
    header_style = ParagraphStyle('header', parent=styles['Normal'],
                                   fontSize=10, textColor=colors.grey)
    normal = styles['Normal']
    normal.fontSize = 11

    elements = []

    # QR code
    qr_bytes = generate_qr_code(qr_code_str)
    qr_img = RLImage(io.BytesIO(qr_bytes), width=3*cm, height=3*cm)

    # Meeting photo — o'ng yuqori burchakda
    photo_img = None
    if meeting.photo:
        try:
            photo_img = RLImage(meeting.photo.path, width=3*cm, height=3.5*cm)
        except Exception:
            pass

    # Sarlavha + rasm — o'ng yuqori
    title_left_style = ParagraphStyle('title_left', parent=styles['Heading1'],
                                       fontSize=16, spaceAfter=4, alignment=0,
                                       textColor=colors.HexColor('#1a5276'))
    sub_left_style = ParagraphStyle('sub_left', parent=styles['Normal'],
                                    fontSize=9, textColor=colors.grey)
    header_content = [
        Paragraph("YOUTHGUARD", title_left_style),
        Paragraph("Yoshlar bilan uchrashuv dalolatnomasi", sub_left_style),
    ]
    from reportlab.platypus import KeepTogether
    header_row = [[header_content, photo_img or '']]
    header_tbl = Table(header_row, colWidths=[13*cm, 4*cm])
    header_tbl.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
        ('LEFTPADDING', (0, 0), (-1, -1), 0),
        ('RIGHTPADDING', (0, 0), (-1, -1), 0),
        ('TOPPADDING', (0, 0), (-1, -1), 0),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
    ]))
    elements.append(header_tbl)
    elements.append(Spacer(1, 0.4*cm))

    # Info table
    data = [
        ['QR kod:', qr_code_str, '', qr_img],
        ['Rahbar:', meeting.rahbar.get_full_name() or meeting.rahbar.username, '', ''],
        ['Yosh:', meeting.youth.full_name, '', ''],
        ['Sana:', meeting.date.strftime('%d.%m.%Y %H:%M'), '', ''],
        ['Holat:', meeting.get_status_display(), '', ''],
        ['Manzil:', meeting.location_address or '—', '', ''],
        ['GPS:', f"{meeting.latitude:.6f}, {meeting.longitude:.6f}" if meeting.latitude else '—', '', ''],
    ]

    ts = [
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('ROWBACKGROUNDS', (0, 0), (1, -1), [colors.HexColor('#f8f9fa'), colors.white]),
        ('GRID', (0, 0), (1, -1), 0.5, colors.HexColor('#dee2e6')),
        ('SPAN', (3, 0), (3, 6)),
        ('VALIGN', (3, 0), (3, 6), 'MIDDLE'),
        ('ALIGN', (3, 0), (3, 6), 'CENTER'),
        ('BACKGROUND', (3, 0), (3, -1), colors.white),
        ('PADDING', (0, 0), (-1, -1), 6),
    ]
    table = Table(data, colWidths=[4*cm, 9*cm, 0.5*cm, 3.5*cm])
    table.setStyle(TableStyle(ts))
    elements.append(table)
    elements.append(Spacer(1, 0.5*cm))

    if meeting.notes:
        import json as _j, re as _r
        _strip_emoji = lambda s: _r.sub(r'[^\x00-\x7FЀ-ӿÀ-ɏ\s]', '', str(s))
        if '[ANSWERS_JSON]' in meeting.notes:
            try:
                survey_answers = _j.loads(meeting.notes.split('[ANSWERS_JSON]')[1].strip())
            except Exception:
                survey_answers = []
            if survey_answers:
                elements.append(Paragraph("<b>So'rovnoma javoblari:</b>", normal))
                elements.append(Spacer(1, 0.15*cm))
                for i, ans in enumerate(survey_answers, 1):
                    q_txt = _strip_emoji(ans.get('q', ''))
                    q_type = ans.get('type', '')
                    val = ans.get('value', '')
                    if q_type == 'photo':
                        val_txt = '[Rasm]'
                    elif q_type == 'location':
                        val_txt = f"GPS: {val}"
                    else:
                        val_txt = _strip_emoji(val)
                    elements.append(Paragraph(f"{i}. <b>{q_txt}:</b> {val_txt}", normal))
                elements.append(Spacer(1, 0.3*cm))
        else:
            clean = _strip_emoji(meeting.notes)
            if clean.strip():
                elements.append(Paragraph(f"<b>Izoh:</b> {clean}", normal))
                elements.append(Spacer(1, 0.3*cm))

    if meeting.impossible_reason:
        elements.append(Paragraph(f"<b>Imkonsizlik sababi:</b> {meeting.impossible_reason}", normal))

    # Footer
    elements.append(Spacer(1, 1*cm))
    footer = Table(
        [['Rahbar imzosi: _____________', 'Yetakchi imzosi: _____________', 'Sana: ___________']],
        colWidths=[6*cm, 6*cm, 5*cm]
    )
    footer.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
    ]))
    elements.append(footer)

    doc.build(elements)
    pdf_bytes = buf.getvalue()

    document, _ = Document.objects.get_or_create(meeting=meeting, defaults={'qr_code': qr_code_str})
    document.file.save(f"meeting_{meeting.id}.pdf", ContentFile(pdf_bytes), save=True)
    document.qr_code = qr_code_str
    document.save()
    return document


def generate_word_doc(meeting, answers=None):
    """Word hujjat — uchrashuv natijalari (on-the-fly)."""
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    import io

    doc = DocxDocument()

    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Sarlavha
    title = doc.add_heading('YOUTHGUARD', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph('Yoshlar bilan uchrashuv dalolatnomasi')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(12)
    sub.runs[0].font.color.rgb = RGBColor(0x6C, 0x75, 0x7D)
    doc.add_paragraph()

    # Asosiy ma'lumotlar jadvali
    tbl = doc.add_table(rows=5, cols=2)
    tbl.style = 'Table Grid'
    rows_data = [
        ('Sana:', meeting.date.strftime('%d.%m.%Y %H:%M')),
        ('Rahbar:', meeting.rahbar.get_full_name() or meeting.rahbar.username),
        ('Yosh:', meeting.youth.full_name),
        ('Holat:', meeting.get_status_display()),
        ('GPS:', f"{meeting.latitude:.6f}, {meeting.longitude:.6f}" if meeting.latitude else '—'),
    ]
    for i, (k, v) in enumerate(rows_data):
        tbl.rows[i].cells[0].text = k
        run = tbl.rows[i].cells[0].paragraphs[0].runs[0]
        run.font.bold = True
        tbl.rows[i].cells[1].text = v

    doc.add_paragraph()

    # So'rovnoma javoblari
    if answers:
        doc.add_heading("So'rovnoma javoblari", level=2)
        for i, ans in enumerate(answers, 1):
            p = doc.add_paragraph(style='List Number')
            run = p.add_run(f"{ans.get('q', '')}: ")
            run.font.bold = True
            if ans.get('type') == 'photo':
                p.add_run('[Rasm]')
            elif ans.get('type') == 'location':
                p.add_run(f"📍 {ans.get('value', '')}")
            else:
                p.add_run(str(ans.get('value', '')))
    elif meeting.notes:
        doc.add_heading('Izohlar', level=2)
        clean = meeting.notes.split('[ANSWERS_JSON]')[0].strip()
        if clean:
            doc.add_paragraph(clean)

    # Rasm
    if meeting.photo:
        try:
            doc.add_paragraph()
            doc.add_heading('Rasm', level=2)
            doc.add_picture(meeting.photo.path, width=Cm(7))
        except Exception:
            pass

    # Imzo qatori
    doc.add_paragraph()
    ftbl = doc.add_table(rows=1, cols=3)
    ftbl.rows[0].cells[0].text = 'Rahbar imzosi: ___________'
    ftbl.rows[0].cells[1].text = 'Yetakchi imzosi: ___________'
    ftbl.rows[0].cells[2].text = f"Sana: {meeting.date.strftime('%d.%m.%Y')}"

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
